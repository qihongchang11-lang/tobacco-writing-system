"""
版式导出Agent
将最终稿件导出为DOCX等格式
"""

from typing import Dict, Any
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import AgentResponse, settings, safe_filename, ensure_file_extension
from .base_agent import BaseAgent

class FormatExporterAgent(BaseAgent):
    """版式导出Agent"""
    
    def __init__(self):
        super().__init__(
            name="FormatExporter",
            description="将最终稿件导出为标准的DOCX文档格式"
        )
        
        self.export_path = settings.data_path / "exports"
        self.export_path.mkdir(parents=True, exist_ok=True)
    
    def get_system_prompt(self) -> str:
        return ""  # 不需要LLM，直接处理
    
    async def process(self, input_data: Dict[str, Any]) -> AgentResponse:
        """处理格式导出"""
        try:
            fact_check_result = input_data.get("fact_check_result")
            
            if not fact_check_result:
                return AgentResponse(
                    success=False,
                    message="缺少事实校对结果",
                    agent_name=self.name
                )
            
            # 生成DOCX文档
            docx_path = self._create_docx_document(fact_check_result.corrected_content)
            
            # 生成Markdown文档
            md_path = self._create_markdown_document(fact_check_result.corrected_content)
            
            return AgentResponse(
                success=True,
                message="格式导出完成",
                data={
                    "docx_path": str(docx_path),
                    "markdown_path": str(md_path),
                    "export_info": {
                        "export_time": datetime.now().isoformat(),
                        "formats": ["DOCX", "Markdown"]
                    }
                },
                agent_name=self.name
            )
            
        except Exception as e:
            return AgentResponse(
                success=False,
                message=f"格式导出失败: {str(e)}",
                agent_name=self.name
            )
    
    def _create_docx_document(self, content: str) -> str:
        """创建DOCX文档"""
        doc = Document()
        
        # 解析内容
        lines = content.split('\n\n')
        title = lines[0].strip() if lines else "无标题"
        body_content = lines[1:] if len(lines) > 1 else []
        
        # 添加标题
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加正文
        for paragraph_text in body_content:
            if paragraph_text.strip():
                paragraph = doc.add_paragraph(paragraph_text.strip())
                paragraph.paragraph_format.first_line_indent = Inches(0.5)
        
        # 添加页脚
        footer_paragraph = doc.add_paragraph(f"\n生成时间：{datetime.now().strftime('%Y年%m月%d日')}")
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 保存文档
        filename = safe_filename(f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        file_path = self.export_path / filename
        
        doc.save(str(file_path))
        self.logger.info(f"DOCX文档已保存: {file_path}")
        
        return str(file_path)
    
    def _create_markdown_document(self, content: str) -> str:
        """创建Markdown文档"""
        lines = content.split('\n\n')
        title = lines[0].strip() if lines else "无标题"
        body_content = lines[1:] if len(lines) > 1 else []
        
        # 构建Markdown内容
        markdown_content = f"# {title}\n\n"
        
        for paragraph_text in body_content:
            if paragraph_text.strip():
                markdown_content += f"{paragraph_text.strip()}\n\n"
        
        markdown_content += f"\n---\n\n*生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}*\n"
        
        # 保存文件
        filename = safe_filename(f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md")
        file_path = self.export_path / filename
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        self.logger.info(f"Markdown文档已保存: {file_path}")
        
        return str(file_path)